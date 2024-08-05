import os
import docx

def get_directory_structure(rootdir):
    """
    Creates a nested dictionary that represents the folder structure of rootdir
    """
    dir_structure = {}
    for dirpath, dirnames, filenames in os.walk(rootdir):
        folder = os.path.relpath(dirpath, rootdir)
        subdir = dir_structure
        if folder != '.':
            for part in folder.split(os.sep):
                subdir = subdir.setdefault(part, {})
        for dirname in dirnames:
            subdir[dirname] = {}
        for filename in filenames:
            subdir[filename] = None
    return dir_structure

def add_structure_to_doc(doc, structure, indent=0):
    """
    Recursively adds folder structure to the doc
    """
    for key, value in structure.items():
        if value is None:
            doc.add_paragraph('  ' * indent + key, style='List Bullet')
        else:
            doc.add_paragraph('  ' * indent + key)
            add_structure_to_doc(doc, value, indent + 1)

# Define the root directory
root_directory = r'C:\Users\scott\Documents\OT System Prototype'  # Update this path to your project directory

# Get the directory structure
directory_structure = get_directory_structure(root_directory)

# Create a new Document
doc = docx.Document()

# Add a title
doc.add_heading('Project Directory Structure', 0)

# Add the directory structure to the document
add_structure_to_doc(doc, directory_structure)

# Save the document
file_path = r"Project_Directory_Structure.docx"  # Update this path where you want to save the document
doc.save(file_path)
print(f"Document saved to {file_path}")
